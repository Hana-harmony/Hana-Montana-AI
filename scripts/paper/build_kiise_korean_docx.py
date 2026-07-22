# 긴 논문 문구와 참고문헌은 원문 보존을 위해 줄 길이 검사에서 제외한다.
# ruff: noqa: E501

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs/paper/acl/k-fnspid-v4-ko.tex"
OUTPUT = ROOT / "output/docx/k-fnspid-kiise-ko.docx"

TITLE_KO = "K-FNSPID: 시간 누수를 통제한 한국 금융 뉴스·공시 데이터셋과 대상 종목 감성·시장영향 분석"
TITLE_EN = "K-FNSPID: A Leakage-Controlled Korean Financial News and Disclosure Dataset for Target-Aware Sentiment and Market-Impact Analysis"
SYSTEM_NAME = "Hana Montana AI(KF-DeBERTa + K-FNSPID)"

CITATIONS = {
    "dong-etal-2024-fnspid": 1,
    "son-etal-2024-krx": 2,
    "son-etal-2025-finkrx": 3,
    "son-etal-2023-korfinasc": 4,
    "tang-etal-2023-finentity": 5,
    "chen-etal-2024-efsa": 6,
    "kim-shin-2022-krfinbert": 7,
    "jeon-etal-2023-kfdeberta": 8,
    "hu-etal-2022-lora": 9,
    "liang-etal-2021-rdrop": 10,
    "cui-etal-2019-classbalanced": 11,
    "guo-etal-2023-predict": 12,
}

REFERENCES = [
    "Z. Dong, X. Fan, and Z. Peng, “FNSPID: A comprehensive financial news dataset in time series,” Proc. 30th ACM SIGKDD, 2024. doi: 10.1145/3637528.3671629.",
    "G. Son, H. Jeon, C. Hwang, and H. Jung, “KRX Bench: Automating financial benchmark creation via large language models,” Proc. Joint Workshop of FinNLP, KDF, and ECONLP, pp. 10–20, 2024.",
    "G. Son, H. Ko, H. Jung, and C. Hwang, “FINKRX: Establishing best practices for Korean financial NLP,” Proc. ACL Industry Track, pp. 1161–1174, 2025. doi: 10.18653/v1/2025.acl-industry.81.",
    "G. Son, H. Lee, N. Kang, and M. Hahm, “Removing non-stationary knowledge from pre-trained language models for entity-level sentiment classification in finance,” AAAI Workshop on Multimodal AI for Financial Forecasting, 2023.",
    "Y. Tang, Y. Yang, A. Huang, A. Tam, and J. Tang, “FinEntity: Entity-level sentiment classification for financial texts,” Proc. EMNLP, pp. 15465–15471, 2023.",
    "T. Chen et al., “EFSA: Towards event-level financial sentiment analysis,” Proc. ACL, pp. 7455–7467, 2024.",
    "E. Kim and H. Shin, “KR-FinBERT: Fine-tuning KR-FinBERT for sentiment analysis,” 2022. https://huggingface.co/snunlp/KR-FinBert-SC.",
    "E. Jeon, J. Kim, M. Song, and J. Ryu, “KF-DeBERTa: Financial domain-specific pre-trained language model,” Proc. HCLT, pp. 143–148, 2023.",
    "E. J. Hu et al., “LoRA: Low-rank adaptation of large language models,” Proc. ICLR, 2022.",
    "X. Liang et al., “R-Drop: Regularized dropout for neural networks,” Advances in Neural Information Processing Systems, vol. 34, pp. 10890–10905, 2021.",
    "Y. Cui, M. Jia, T.-Y. Lin, Y. Song, and S. Belongie, “Class-balanced loss based on effective number of samples,” Proc. IEEE/CVF CVPR, pp. 9268–9277, 2019.",
    "Y. Guo, C. Hu, and Y. Yang, “Predict the future from the past? On the temporal data distribution shift in financial sentiment classifications,” Proc. EMNLP, pp. 1029–1038, 2023.",
]


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    fonts = run._element.rPr.rFonts
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:cs"), name)
    fonts.set(qn("w:hint"), "eastAsia")
    lang = OxmlElement("w:lang")
    lang.set(qn("w:eastAsia"), "ko-KR")
    run._element.rPr.append(lang)


def set_cell_margins(cell, top: int = 50, start: int = 70, bottom: int = 50, end: int = 70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), value)


def set_section_geometry(section, columns: int = 1) -> None:
    section.page_width = Mm(182)
    section.page_height = Mm(257)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(15)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), "454")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char, instr, end))


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "AppleGothic"
    normal.font.size = Pt(9.2)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleGothic")
    normal._element.rPr.rFonts.set(qn("w:cs"), "AppleGothic")
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.first_line_indent = Mm(3)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in (
        ("Heading 1", 13, 8, 3),
        ("Heading 2", 10.5, 5, 2),
    ):
        style = styles[style_name]
        style.font.name = "AppleGothic"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleGothic")
        style._element.rPr.rFonts.set(qn("w:cs"), "AppleGothic")
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Reference" not in styles:
        ref_style = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref_style = styles["Reference"]
    ref_style.font.name = "AppleGothic"
    ref_style.font.size = Pt(7.5)
    ref_style._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleGothic")
    ref_style._element.rPr.rFonts.set(qn("w:cs"), "AppleGothic")
    ref_style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    ref_style.paragraph_format.left_indent = Mm(5)
    ref_style.paragraph_format.first_line_indent = Mm(-5)
    ref_style.paragraph_format.line_spacing = 1.0
    ref_style.paragraph_format.space_after = Pt(1)


def add_labeled_paragraph(doc: Document, label: str, body: str, english: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_run_font(r, "Times New Roman" if english else "AppleGothic", 9.2, True)
    r = p.add_run("  " + body)
    set_run_font(r, "Times New Roman" if english else "AppleGothic", 9.2)


def clean_latex(text: str) -> str:
    text = text.replace("\\dataset{}", "K-FNSPID").replace("\\system{}", SYSTEM_NAME)
    text = text.replace("\\%", "%").replace("~", " ")
    text = text.replace("quadratic weighted kappa", "quadratic weighted kappa")
    text = re.sub(r"\\textbf\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\ref\{tab:sent-ko\}", "2", text)
    text = re.sub(r"\\ref\{tab:impact-ko\}", "3", text)

    def replace_cite(match: re.Match[str]) -> str:
        numbers = sorted(CITATIONS[key] for key in match.group(1).split(",") if key in CITATIONS)
        return "[" + ", ".join(str(number) for number in numbers) + "]"

    text = re.sub(r"\\cite\{([^{}]+)\}", replace_cite, text)
    text = re.sub(r"\\[A-Za-z]+\{([^{}]*)\}", r"\1", text)
    text = text.replace("\\&", "&").replace("--", "–")
    return re.sub(r"\s+", " ", text).strip()


def add_table(doc: Document, label: str) -> None:
    if label == "tab:dataset-ko":
        rows = [
            ("구성요소", "건수"),
            ("뉴스 문서", "524,696"),
            ("OpenDART 공시", "722,989"),
            ("전체 문서", "1,247,685"),
            ("문서-종목 관계", "1,136,118"),
            ("시장반응 표본", "715,015"),
            ("비혼입 대표 표본", "255,168"),
            ("일별 시세", "10,691,998"),
            ("감성 학습 Gold", "1,794"),
            ("감성 개발 Gold", "895"),
            ("감성 Silver", "32,700"),
        ]
        caption = "표 1. K-FNSPID의 구성과 규모"
        widths = (Mm(47), Mm(25))
    elif label == "tab:sent-ko":
        rows = [
            ("모델", "국내 뉴스", "국내 공시"),
            ("제안 모델", "0.7503 / 0.5530", "0.8646 / 0.6024"),
            ("원본 KR-FinBERT-SC", "0.5781 / 0.4937", "0.8535 / 0.6146"),
            ("공정 full-FT", "0.7677 / 0.5771", "0.8514 / 0.5647"),
        ]
        caption = "표 2. 잠금 후 확증 Gold의 설계가중 Accuracy / Macro-F1 비교"
        widths = (Mm(31), Mm(20.5), Mm(20.5))
    else:
        rows = [
            ("출처", "모델", "M-F1", "QWK"),
            ("국내 뉴스", "KR-FinBERT-SC", "0.3506", "0.3962"),
            ("", "KF-DeBERTa", "0.3745", "0.4754"),
            ("국내 공시", "KR-FinBERT-SC", "0.3131", "0.1611"),
            ("", "KF-DeBERTa", "0.3216", "0.1550"),
        ]
        caption = "표 3. 시간 분할 Test의 시장영향 중요도 성능"
        widths = (Mm(15), Mm(27), Mm(15), Mm(15))

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_width_twips = sum(int(width) for width in widths) // 635
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(table_width_twips))
    table_width.set(qn("w:type"), "dxa")
    for col_index, width in enumerate(widths):
        table.columns[col_index].width = width
        table._tbl.tblGrid.gridCol_lst[col_index].set(qn("w:w"), str(int(width) // 635))
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.width = widths[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Mm(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or col_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, "AppleGothic", 7.4, row_index == 0)
            borders: dict[str, dict[str, str]] = {}
            if row_index == 0:
                borders["top"] = {"val": "single", "sz": "8", "color": "000000"}
                borders["bottom"] = {"val": "single", "sz": "6", "color": "000000"}
            if row_index == len(rows) - 1:
                borders["bottom"] = {"val": "single", "sz": "8", "color": "000000"}
            if borders:
                set_cell_border(cell, **borders)
    caption_p = doc.add_paragraph()
    caption_p.paragraph_format.first_line_indent = Mm(0)
    caption_p.paragraph_format.space_before = Pt(2)
    caption_p.paragraph_format.space_after = Pt(4)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption_p.add_run(caption)
    set_run_font(r, "AppleGothic", 8)


def add_body_from_latex(doc: Document) -> None:
    source = SOURCE.read_text(encoding="utf-8")
    body = source.split("\\section{서론}", 1)[1]
    body = "\\section{서론}\n" + body.split("\\begingroup", 1)[0]
    lines = body.splitlines()
    paragraph: list[str] = []
    in_table = False
    table_lines: list[str] = []
    section_number = 0
    subsection_number = 0

    def flush_paragraph() -> None:
        if not paragraph:
            return
        text = clean_latex(" ".join(paragraph))
        paragraph.clear()
        if not text:
            return
        p = doc.add_paragraph(text)
        p.paragraph_format.keep_together = False

    for raw in lines:
        line = raw.strip()
        if line.startswith("\\begin{table}"):
            flush_paragraph()
            in_table = True
            table_lines = [line]
            continue
        if in_table:
            table_lines.append(line)
            if line.startswith("\\end{table}"):
                match = re.search(r"\\label\{([^{}]+)\}", " ".join(table_lines))
                if match:
                    add_table(doc, match.group(1))
                in_table = False
            continue
        section = re.match(r"\\section\{(.+)\}", line)
        subsection = re.match(r"\\subsection\{(.+)\}", line)
        if section:
            flush_paragraph()
            section_number += 1
            subsection_number = 0
            doc.add_paragraph(f"{section_number}. {section.group(1)}", style="Heading 1")
        elif subsection:
            flush_paragraph()
            subsection_number += 1
            doc.add_paragraph(
                f"{section_number}.{subsection_number} {subsection.group(1)}",
                style="Heading 2",
            )
        elif not line:
            flush_paragraph()
        elif line.startswith("\\") and line.split("{", 1)[0] in {
            "\\begin",
            "\\end",
            "\\centering",
            "\\small",
            "\\resizebox",
            "\\setlength",
        }:
            continue
        else:
            paragraph.append(line)
    flush_paragraph()


def build() -> None:
    doc = Document()
    configure_styles(doc)
    set_section_geometry(doc.sections[0], 1)
    doc.sections[0].different_first_page_header_footer = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE_KO)
    set_run_font(r, "AppleGothic", 16, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE_EN)
    set_run_font(r, "Times New Roman", 11.5)

    for text, font, size, bold, after in (
        ("최성현", "AppleGothic", 11.5, True, 0),
        ("Sunghyun Choi", "Times New Roman", 10, False, 3),
        ("한국공학대학교 컴퓨터공학부 소프트웨어학과", "AppleGothic", 9.2, False, 0),
        ("Department of Software, School of Computer Engineering, Tech University of Korea", "Times New Roman", 8.7, False, 14),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, font, size, bold)

    summary = (
        "본 연구는 한국 상장기업 뉴스 524,696건과 전자공시 722,989건을 일별 시세 10,691,998행에 연결한 K-FNSPID를 제안한다. "
        "데이터셋은 문서 1,247,685건, 문서-종목 관계 1,136,118건과 사건 중복 및 동일 거래일 혼입을 제거한 시장반응 표본 255,168건으로 구성된다. "
        "각 문서는 발표 시각을 한국 거래일에 맞추고, 반복 보도와 복수 사건의 영향을 통제하도록 정규화하였다. 감성분류에는 대상 종목을 입력에 명시한 KF-DeBERTa 기반 분류기를 사용하고, 뉴스와 공시의 표현 차이를 반영하는 출처별 잔차 계층을 결합하였다. "
        "잠금 후 600건씩 평가한 감성분류의 설계가중 Accuracy/Macro-F1은 국내 뉴스 0.7503/0.5530, 국내 공시 0.8646/0.6024였다. 두 출처의 공동 우월성이 확인되지 않아 후보를 배포 모델로 승격하지 않았다. "
        "동일 K-FNSPID 시간 Test에서 시장영향 중요도 Macro-F1은 국내 뉴스 0.3745, 국내 공시 0.3216으로 KR-FinBERT-SC보다 각각 6.82%(0.0239점), 2.72%(0.0085점) 높았다. "
        "거래일 군집 검정에서 국내 뉴스의 우위는 확인했지만 국내 공시는 확인하지 못했다. 본 연구는 한국 금융 문서, 대상 종목과 시세를 시간 순서에 따라 연결하고 출처별 성능과 불확실성을 분리해 보고하는 재현 가능한 평가 기반을 제공한다."
    )
    abstract = (
        "This study presents K-FNSPID, a Korean financial dataset linking 524,696 news articles and 722,989 regulatory disclosures to 10,691,998 daily price records. "
        "The release contains 1,247,685 documents, 1,136,118 document-security relations, and 255,168 market-response samples after controlling for duplicate events and same-day event collisions. "
        "Publication times are mapped to Korean trading sessions, and repeated reports are grouped at the event level. For target-aware sentiment classification, we employ a KF-DeBERTa classifier that explicitly conditions on the target security and models source-specific residuals for news and disclosures. "
        "In the locked 600-item-per-source sentiment evaluation, design-weighted accuracy/macro-F1 is 0.7503/0.5530 for news and 0.8646/0.6024 for disclosures. Joint source-wise superiority is not established, so the candidate is not promoted. "
        "For ordinal market-impact classification, the model obtains macro-F1 scores of 0.3745 on news and 0.3216 on disclosures, compared with 0.3506 and 0.3131 for KR-FinBERT-SC under the same K-FNSPID temporal tests. "
        "Clustered inference confirms superiority for news but not for disclosures. The resulting resource provides a reproducible basis for Korean financial language evaluation with explicit temporal and source controls."
    )
    add_labeled_paragraph(doc, "요약", summary)
    add_labeled_paragraph(doc, "키워드:", "금융 자연어처리, 금융 감성분석, 전자공시, 시장영향 분류, 시간 누수")
    add_labeled_paragraph(doc, "Abstract", abstract, english=True)
    add_labeled_paragraph(doc, "Keywords:", "financial NLP, financial sentiment analysis, regulatory disclosure, market-impact classification, temporal leakage", english=True)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(body_section, 2)
    body_section.header.is_linked_to_previous = False
    header = body_section.header.paragraphs[0]
    header.paragraph_format.first_line_indent = Mm(0)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("K-FNSPID: 한국 금융 뉴스·공시 데이터셋")
    set_run_font(r, "AppleGothic", 8)
    tab_stops = header.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Mm(150))
    header.add_run("\t")
    add_page_number(header)

    add_body_from_latex(doc)
    doc.add_paragraph("참고문헌", style="Heading 1")
    for index, reference in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph(style="Reference")
        p.add_run(f"[{index}] {reference}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
